#!/usr/bin/env python3
"""
Generates Omid Zanganeh - Resume.docx
Professional, ATS-safe Word resume with clean modern design.
Run: python build_resume.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

# ─── Palette ──────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x3A, 0x5C)   # deep navy — headings / name
TEAL      = RGBColor(0x0D, 0x7A, 0x8A)   # accent teal — section rules / pills
TEAL_LIGHT = RGBColor(0xE4, 0xF4, 0xF6)  # very light teal — section header bg
GRAY_MID  = RGBColor(0x55, 0x65, 0x75)   # mid gray — company / dates
GRAY_DIM  = RGBColor(0x88, 0x96, 0xA3)   # dim gray — location / labels
BLACK     = RGBColor(0x1A, 0x1A, 0x1A)   # near-black — body text
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Font ─────────────────────────────────────────────────────────────────────
BODY_FONT = "Calibri"


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def no_border_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def para_border_bottom(para, color='0D7A8A', size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run(para, text, bold=False, italic=False, size=10,
            color=None, font=BODY_FONT, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc_or_cell, text='', align=WD_ALIGN_PARAGRAPH.LEFT,
             before=0, after=0):
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before=before, after=after)
    return p


def section_header(doc, title: str):
    """Teal underline section header — ATS readable, visually clean."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=50, after=10)
    para_border_bottom(p)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.font.all_caps = True
    # letter spacing via rPr
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '30')
    rPr.append(spacing)


def bullet_para(doc, text: str, indent=Inches(0.15)):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before=0, after=10)
    # Remove default style indent — set our own
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '216')
    ind.set(qn('w:hanging'), '216')
    pPr.append(ind)
    # Custom bullet char
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)
    # Plain text bullet — List Bullet style provides the bullet character
    run_text = p.add_run(text)
    run_text.font.name = BODY_FONT
    run_text.font.size = Pt(9)
    run_text.font.color.rgb = BLACK
    return p


def job_header(doc, company: str, title: str, dates: str, location: str,
               badge: str = None):
    # Company + dates row
    p = doc.add_paragraph()
    set_para_spacing(p, before=60, after=0)
    add_run(p, company, bold=True, size=10, color=NAVY)
    if badge:
        add_run(p, f'  ✦ {badge}', bold=False, size=8, color=TEAL, italic=True)
    # right-align dates via tab stop
    add_run(p, '\t', size=10)
    add_run(p, dates, size=9, color=GRAY_MID)
    # Tab stop at right margin
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9072')   # ~6.3 in from left
    tabs.append(tab)
    pPr.append(tabs)

    # Title row
    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=2)
    add_run(p2, title, bold=False, italic=True, size=9.5, color=GRAY_MID)

    # Location row
    p3 = doc.add_paragraph()
    set_para_spacing(p3, before=0, after=14)
    add_run(p3, location, size=9, color=GRAY_DIM)


def skill_chips_para(doc, label: str, items: list):
    p = doc.add_paragraph()
    set_para_spacing(p, before=12, after=12)
    add_run(p, label + '  ', bold=True, size=9, color=NAVY)
    add_run(p, '  ·  '.join(items), size=9, color=BLACK)


# ─── Build Document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Zero out default paragraph spacing so our explicit values are exact
    normal = doc.styles['Normal']
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    # US Letter, compact margins for 2-page fit
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.top_margin    = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin   = Inches(0.55)
        section.right_margin  = Inches(0.55)

    # ── HEADER ────────────────────────────────────────────────────────────────
    # Name
    name_p = doc.add_paragraph()
    set_para_spacing(name_p, before=0, after=10)
    name_run = name_p.add_run('OMID ZANGANEH')
    name_run.font.name = BODY_FONT
    name_run.font.size = Pt(24)
    name_run.bold = True
    name_run.font.color.rgb = NAVY

    # Title line
    title_p = doc.add_paragraph()
    set_para_spacing(title_p, before=0, after=2)
    add_run(title_p, 'Senior GIS Developer and Analyst', bold=True, size=12, color=NAVY)

    # Degree line
    degree_p = doc.add_paragraph()
    set_para_spacing(degree_p, before=0, after=30)
    add_run(degree_p, 'MS Geography — Geographic Information Science & Technology  ·  Workflow Automation  ·  AI/ML Integration',
            size=10, color=TEAL)

    # Contact bar — single paragraph with separators
    contact_p = doc.add_paragraph()
    set_para_spacing(contact_p, before=0, after=0)
    para_border_bottom(contact_p, color='0D7A8A', size=6)
    contacts = [
        '+1 (531) 229-6873',
        'ozanganeh@unomaha.edu',
        'linkedin.com/in/omidzanganeh',
        'omidzanganeh.com',
        'Lincoln, Nebraska',
    ]
    for i, c in enumerate(contacts):
        add_run(contact_p, c, size=9, color=GRAY_MID)
        if i < len(contacts) - 1:
            add_run(contact_p, '   |   ', size=9, color=TEAL)

    # ── PROFESSIONAL SUMMARY ──────────────────────────────────────────────────
    section_header(doc, 'Professional Summary')
    summary_p = doc.add_paragraph()
    set_para_spacing(summary_p, before=10, after=0)
    add_run(summary_p,
        'GIS Developer and software engineer specializing in GIS workflow automation, '
        'ArcGIS Pro development, and AI-powered spatial workflows. Builds production desktop '
        'applications (Python, C#), ArcGIS Pro add-ins and geoprocessing toolboxes, SQL Server '
        'data pipelines, and full-stack web tools (Next.js, React, TypeScript). Known for '
        'compressing multi-day manual processes into automated pipelines. 2026 Edison Award winner '
        'for demonstrating superior technical ability. MS Geography (GIS&T), 4.0 GPA.',
        size=9.5, color=BLACK)

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────────
    section_header(doc, 'Work Experience')

    # Olsson
    job_header(doc,
               company='Olsson',
               title='GIS Developer',
               dates='March 2025 – Present',
               location='Lincoln, Nebraska  ✦  2025 Nominee & 2026 Edison Award Winner')

    olsson_bullets = [
        'Architected production Python and C# desktop applications for GIS workflows including '
        'automated bore profile generation that cut processing time from days to minutes.',

        'Built an ArcGIS Automation Suite of custom Python geoprocessing toolboxes reducing '
        'manual GIS steps by ~90% and accelerating fiber network design timelines.',

        'Developed ArcGIS Pro add-ins (.NET/C#/WPF): multi-source GIS Data Downloader '
        '(OSM, USGS, FEMA, Census/TIGER, BSL), RF Analysis panel, Street View map tool, '
        'and FTTH network design dock pane.',

        'Engineered AI-powered tools using Azure AI Foundry, Azure OpenAI, and Google AI Studio, '
        'including RFP Radar for intelligent web-grounded contract search and classification, '
        'cutting strategic sourcing timelines from months to hours.',

        'Built AI object detection applications for remote inventory of utility poles, '
        'streetlights, and telecom infrastructure.',

    ]
    for b in olsson_bullets:
        bullet_para(doc, b)

    # UNO Teaching
    job_header(doc,
               company='University of Nebraska at Omaha',
               title='Graduate Teaching Assistant — Instructor of Record',
               dates='January 2024 – August 2025',
               location='Omaha, Nebraska')
    bullet_para(doc, 'Taught Human-Environment Geography lab sections as sole instructor of record '
                     'to 150+ students across three semesters, emphasizing hands-on spatial analysis '
                     'and real-world applications.')

    # Spatial Justice
    job_header(doc,
               company='University of Nebraska at Omaha',
               title='GIS Technician — Omaha Spatial Justice Project',
               dates='June 2024 – August 2025',
               location='Omaha, Nebraska')
    bullet_para(doc, 'Digitized historical land parcels from archival documents and aerial photography; '
                     'reviewed legal records to build an accurate geodatabase of racially restrictive '
                     'covenants in Douglas County, supporting urban spatial justice research.')

    # ── SELECTED PROJECTS ─────────────────────────────────────────────────────
    section_header(doc, 'Selected Projects')

    projects = [
        ('RFP Radar', 'AI-Powered RFP Sourcing',
         'Azure OpenAI · Google Gemini · Playwright · Python/C#',
         'Production desktop app for intelligent, web-grounded RFP search, classification, '
         'batch county search, and CSV export. Reduced sourcing time from months to hours.'),
        ('Bore Profile Automation', 'Directional Drilling Profile Generator',
         'Python · C# · ArcGIS Pro · SQL Server · Matplotlib',
         'Fully automated app that reads spatial waypoints, processes elevation models, '
         'and generates 2D/3D bore profiles. Cut processing from days to minutes.'),
        ('ArcGIS Data Downloader Add-in', 'Multi-Source GIS Data Acquisition',
         'C# · .NET 8 · ArcGIS Pro SDK · WPF · REST APIs',
         'ArcGIS Pro dock-pane add-in that downloads OSM, USGS, FEMA, Census/TIGER, '
         'Wikipedia, and BSL layers directly into projects.'),
        ('Fiber Automatic Expansion', 'Automated Fiber Build-Area Planning',
         'C# · .NET · ArcGIS Pro SDK · WPF · Spatial Analysis',
         'ArcGIS Pro decision-support add-in that grids the study area for density '
         'analysis, flood-fills viable build zones against configurable PPM thresholds, '
         'and generates economic scorecards with live human-in-the-loop refinement. '
         'Cut expansion planning from days/weeks to minutes.'),
        ('FTTH Network Designer', 'Automated Fiber Optic Network Planning',
         'C# · .NET 8 · ArcGIS Pro SDK · WPF · Kruskal MST',
         'ArcGIS Pro add-in automating FTTH layout in three steps: place shafts at road '
         'intersections, connect homes to nearest shaft (tagging Street vs Driveway Drop by '
         'geodesic length), and run Kruskal\'s MST for an optimal Main Trunk / Terminal '
         'Branch backbone. Reduced network layout from days to minutes.'),
        ('RF Analysis Tool', '8-Tool Wireless Planning Panel',
         'C# · ArcGIS Pro SDK · Python',
         'ArcGIS Pro side panel with coverage prediction, PCI/RSI planner, tilt/azimuth '
         'optimizers, interference analysis, and tower placement optimizer — all as map layers.'),
        ('Aerial & Streetview AI Detection', 'YOLO Utility Infrastructure Detection',
         'Python · YOLO · OpenCV · Aerial/Street Imagery APIs',
         'Desktop apps fetching aerial tiles or traversing street routes, running custom YOLO '
         'models to detect/classify utility assets and export georeferenced results to ArcGIS.'),
        ('GeoPipe', 'Enterprise GIS ETL to SQL Server',
         'Python · CustomTkinter · pyodbc · SQL Server · PyInstaller',
         'GUI ETL tool for large spatial/tabular imports into SQL Server with schema '
         'auto-detection, GEOMETRY/GEOGRAPHY support, and connection-loss auto-resume.'),
        ('omidzanganeh.com', 'Full-Stack Portfolio & Browser GIS Tools',
         'Next.js 16 · React · TypeScript · Leaflet · Vercel',
         'Personal site with 10+ browser GIS tools and general web apps (geocoder, isochrone, '
         'elevation profile, census, coordinate converter, Gymflow, Stock Screener) '
         'and an AI news aggregator.'),
    ]

    for name, subtitle, tech, desc in projects:
        p = doc.add_paragraph()
        set_para_spacing(p, before=36, after=0)
        add_run(p, name, bold=True, size=9.5, color=NAVY)
        add_run(p, f'  —  {subtitle}', size=9, color=TEAL)

        p2 = doc.add_paragraph()
        set_para_spacing(p2, before=0, after=0)
        add_run(p2, tech, size=8.5, color=GRAY_DIM, italic=True)

        p3 = doc.add_paragraph()
        set_para_spacing(p3, before=0, after=8)
        add_run(p3, desc, size=9, color=BLACK)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    section_header(doc, 'Education')

    # MS
    p = doc.add_paragraph()
    set_para_spacing(p, before=50, after=0)
    add_run(p, 'Master of Science, Geography — Geographic Information Science & Technology',
            bold=True, size=10, color=NAVY)
    add_run(p, '\t', size=10)
    add_run(p, 'August 2025', size=9, color=GRAY_MID)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)

    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=0)
    add_run(p2, 'University of Nebraska at Omaha, Nebraska', size=9.5, color=GRAY_MID)
    add_run(p2, '   GPA: 4.00  ·  GRACA Project Award', bold=True, size=9.5, color=TEAL)

    p3 = doc.add_paragraph()
    set_para_spacing(p3, before=4, after=16)
    add_run(p3, 'Thesis: ', bold=True, size=9, color=NAVY)
    add_run(p3,
        'Spatiotemporal Analysis of NOx Emissions from U.S. Cement Plants Using TROPOMI Data '
        '— remote sensing, hotspot analysis, environmental visualization, population exposure '
        '& environmental justice.',
        size=9, color=BLACK)

    # BS
    p = doc.add_paragraph()
    set_para_spacing(p, before=28, after=0)
    add_run(p, 'Bachelor of Science, Geomatics (Surveying) Engineering',
            bold=True, size=10, color=NAVY)
    add_run(p, '\t', size=10)
    add_run(p, 'August 2016', size=9, color=GRAY_MID)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)

    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=40)
    add_run(p2, 'Geomatics College of National Cartographic Center (GCNCC), Tehran',
            size=9.5, color=GRAY_MID)
    set_para_spacing(p2, before=0, after=14)

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────────
    section_header(doc, 'Technical Skills')

    skill_groups = [
        ('Languages & Frameworks',
         ['Python', 'C# / .NET', 'SQL', 'TypeScript', 'JavaScript', 'HTML/CSS',
          'React', 'Next.js', 'WPF']),
        ('GIS & Spatial',
         ['ArcGIS Pro', 'ArcGIS Pro SDK', 'ArcPy', 'Python Toolboxes (.pyt)',
          'ArcGIS Online', 'ArcGIS Enterprise', 'QGIS', 'Geoprocessing',
          'Network Analysis', 'Remote Sensing', 'Google Earth Engine']),
        ('AI / ML',
         ['Azure OpenAI', 'Azure AI Foundry', 'Google AI Studio / Gemini',
          'YOLO', 'OpenCV', 'Prompt Engineering', 'Web Grounding',
          'Batch Classification Pipelines']),
        ('Data & Backend',
         ['SQL Server', 'Supabase', 'PostgreSQL', 'pyodbc',
          'ETL Pipelines', 'Spatial Types (GEOMETRY/GEOGRAPHY)', 'Smartsheet API']),
        ('Web & Cloud',
         ['Next.js App Router', 'Vite', 'Leaflet', 'Vercel',
          'REST APIs', 'PWA', 'Microsoft Azure', 'Google Cloud']),
        ('Desktop & Tooling',
         ['CustomTkinter', 'PyInstaller', 'Playwright', 'Matplotlib',
          'AutoCAD', 'ENVI', 'SNAP', 'Photomod', 'Tableau',
          'Adobe Photoshop', 'Adobe Illustrator']),
        ('Languages',
         ['English (Fluent)', 'Persian (Native)']),
    ]

    for label, items in skill_groups:
        skill_chips_para(doc, label, items)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=20, after=0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_border_bottom(p2, color='CCCCCC', size=4)
    add_run(p2,
        'For project write-ups, live GIS tools, and demos: ',
        size=8.5, color=GRAY_DIM)
    add_run(p2, 'omidzanganeh.com', bold=True, size=8.5, color=TEAL)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    out_path = os.path.join(os.path.dirname(__file__), 'Omid Zanganeh - Resume v2.docx')
    doc.save(out_path)
    print(f'✅ Saved: {out_path}')
    return out_path


if __name__ == '__main__':
    build()
